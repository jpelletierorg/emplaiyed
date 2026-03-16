"""Tests for DOCX rendering of CV and letter."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from emplaiyed.core.models import Address, Profile
from emplaiyed.generation.cv_generator import (
    CVCertification,
    CVEducation,
    CVExperience,
    CVProject,
    GeneratedCV,
    SkillCategory,
)
from emplaiyed.generation.letter_generator import GeneratedLetter
from emplaiyed.rendering.docx_renderer import render_cv_docx, render_letter_docx


def _sample_cv() -> GeneratedCV:
    return GeneratedCV(
        name="Alice Test",
        email="alice@example.com",
        phone="+1-555-0000",
        location="Montreal, QC",
        professional_title="Senior Cloud Architect",
        summary="Seasoned cloud architect with 10+ years of experience.",
        skill_categories=[
            SkillCategory(
                category="Cloud & Infrastructure", skills=["AWS", "GCP", "Terraform"]
            ),
            SkillCategory(category="Programming", skills=["Python", "Go"]),
        ],
        experience=[
            CVExperience(
                company="Acme Corp",
                title="Lead Developer",
                start_date="2020-01",
                end_date="Present",
                description="Led cloud migration.",
                highlights=["Reduced costs by 40%", "Migrated 12 services"],
            ),
        ],
        education=[
            CVEducation(
                institution="Universite Laval",
                degree="M.Sc.",
                field="Computer Science",
                start_date="2010",
                end_date="2012",
            ),
        ],
        certifications=[
            CVCertification(name="AWS SAA", issuer="Amazon", date="2023"),
        ],
        projects=[
            CVProject(
                name="CloudOptimizer",
                description="Open-source cost optimization tool",
                technologies=["Python", "AWS"],
                url="https://github.com/alice/cloudoptimizer",
            ),
        ],
        languages=["English (Native)", "French (Fluent)"],
    )


def _sample_letter() -> GeneratedLetter:
    return GeneratedLetter(
        greeting="Dear Hiring Manager,",
        hook="Your team's challenge with scaling cloud infrastructure resonates with me.",
        proof="At Acme Corp, I reduced infrastructure costs by 40% while improving uptime to 99.99%.",
        close="I'd welcome the chance to discuss how I can help your team achieve similar results.",
        closing="Sincerely,",
        signature_name="Alice Test",
    )


def _sample_profile() -> Profile:
    return Profile(
        name="Alice Test",
        email="alice@example.com",
        phone="+1-555-0000",
        address=Address(city="Montreal", province_state="QC"),
    )


# ---------------------------------------------------------------------------
# CV DOCX tests
# ---------------------------------------------------------------------------


class TestRenderCvDocx:
    def test_creates_file(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_file_is_valid_docx(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        # Should have paragraphs
        assert len(doc.paragraphs) > 0

    def test_contains_name(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Alice Test" in all_text

    def test_contains_skills(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AWS" in all_text
        assert "Python" in all_text

    def test_contains_experience(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in all_text
        assert "Reduced costs by 40%" in all_text

    def test_contains_projects(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "CloudOptimizer" in all_text

    def test_contains_education(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Universite Laval" in all_text

    def test_contains_certifications(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AWS SAA" in all_text

    def test_contains_languages(self, tmp_path: Path) -> None:
        out = tmp_path / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "English" in all_text
        assert "French" in all_text

    def test_creates_parent_dirs(self, tmp_path: Path) -> None:
        out = tmp_path / "nested" / "dir" / "cv.docx"
        render_cv_docx(_sample_cv(), out)
        assert out.exists()

    def test_minimal_cv(self, tmp_path: Path) -> None:
        """A CV with only required fields should render without error."""
        cv = GeneratedCV(
            name="Bob",
            email="bob@example.com",
            professional_title="Developer",
            summary="",
            skill_categories=[],
            experience=[],
            education=[],
        )
        out = tmp_path / "minimal.docx"
        render_cv_docx(cv, out)
        assert out.exists()


# ---------------------------------------------------------------------------
# Letter DOCX tests
# ---------------------------------------------------------------------------


class TestRenderLetterDocx:
    def test_creates_file(self, tmp_path: Path) -> None:
        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_file_is_valid_docx(self, tmp_path: Path) -> None:
        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out)
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0

    def test_contains_greeting(self, tmp_path: Path) -> None:
        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Dear Hiring Manager" in all_text

    def test_contains_body_paragraphs(self, tmp_path: Path) -> None:
        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "scaling cloud infrastructure" in all_text
        assert "reduced infrastructure costs" in all_text

    def test_contains_closing(self, tmp_path: Path) -> None:
        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Sincerely" in all_text
        assert "Alice Test" in all_text

    def test_with_profile_includes_contact(self, tmp_path: Path) -> None:
        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out, profile=_sample_profile())
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "alice@example.com" in all_text
        assert "Montreal" in all_text

    def test_without_profile(self, tmp_path: Path) -> None:
        """Should render fine without a profile."""
        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out, profile=None)
        assert out.exists()

    def test_contains_date(self, tmp_path: Path) -> None:
        """Letter should include today's date."""
        from datetime import datetime

        out = tmp_path / "letter.docx"
        render_letter_docx(_sample_letter(), out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Check that current year appears (robust enough for any test date)
        assert str(datetime.now().year) in all_text

    def test_creates_parent_dirs(self, tmp_path: Path) -> None:
        out = tmp_path / "nested" / "dir" / "letter.docx"
        render_letter_docx(_sample_letter(), out)
        assert out.exists()
