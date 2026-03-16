"""DOCX rendering — ATS-optimized Word documents for CVs and letters."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from emplaiyed.core.models import Profile
from emplaiyed.generation.cv_generator import GeneratedCV
from emplaiyed.generation.letter_generator import GeneratedLetter

_DARK_BLUE = RGBColor(0x1A, 0x36, 0x5D)
_BODY_COLOR = RGBColor(0x2D, 0x37, 0x48)
_MUTED_COLOR = RGBColor(0x71, 0x80, 0x96)


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a section heading in the standard style."""
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = _DARK_BLUE
        run.font.size = Pt(11)


def render_cv_docx(cv: GeneratedCV, path: Path) -> None:
    """Render a GeneratedCV to an ATS-optimized DOCX file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = _BODY_COLOR

    # Name
    name_para = doc.add_heading(cv.name, level=1)
    for run in name_para.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = _DARK_BLUE

    # Professional title
    title_para = doc.add_paragraph(cv.professional_title)
    title_para.style = doc.styles["Normal"]
    for run in title_para.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Contact info (single line, pipe-separated)
    contact_parts = [cv.email]
    if cv.phone:
        contact_parts.append(cv.phone)
    if cv.location:
        contact_parts.append(cv.location)
    if cv.linkedin:
        contact_parts.append(cv.linkedin)
    if cv.github:
        contact_parts.append(cv.github)
    contact_para = doc.add_paragraph(" | ".join(contact_parts))
    for run in contact_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = _MUTED_COLOR

    # Summary
    if cv.summary:
        _add_section_heading(doc, "Summary")
        doc.add_paragraph(cv.summary)

    # Skills
    if cv.skill_categories:
        _add_section_heading(doc, "Skills")
        for cat in cv.skill_categories:
            p = doc.add_paragraph()
            run = p.add_run(f"{cat.category}: ")
            run.bold = True
            run.font.size = Pt(9)
            p.add_run(", ".join(cat.skills)).font.size = Pt(9)

    # Experience
    if cv.experience:
        _add_section_heading(doc, "Experience")
        for exp in cv.experience:
            # Title — Company | Dates
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.title} — {exp.company}")
            run.bold = True
            run.font.size = Pt(10)
            if exp.start_date or exp.end_date:
                dates = f" | {exp.start_date or '?'} – {exp.end_date or 'Present'}"
                date_run = p.add_run(dates)
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = _MUTED_COLOR
            if exp.description:
                desc_para = doc.add_paragraph(exp.description)
                desc_para.style = doc.styles["Normal"]
                for r in desc_para.runs:
                    r.italic = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = _MUTED_COLOR
            for h in exp.highlights:
                bp = doc.add_paragraph(h, style="List Bullet")
                for r in bp.runs:
                    r.font.size = Pt(9)

    # Projects
    if cv.projects:
        _add_section_heading(doc, "Projects")
        for proj in cv.projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.name)
            run.bold = True
            run.font.size = Pt(10)
            if proj.url:
                p.add_run(f" — {proj.url}").font.size = Pt(9)
            if proj.description:
                doc.add_paragraph(proj.description).style = doc.styles["Normal"]
            if proj.technologies:
                tech_para = doc.add_paragraph(
                    f"Technologies: {', '.join(proj.technologies)}"
                )
                for r in tech_para.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = _MUTED_COLOR

    # Education
    if cv.education:
        _add_section_heading(doc, "Education")
        for edu in cv.education:
            dates = ""
            if edu.start_date or edu.end_date:
                dates = f" ({edu.start_date or '?'} – {edu.end_date or 'Present'})"
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.degree} in {edu.field}")
            run.bold = True
            run.font.size = Pt(9)
            p.add_run(f" — {edu.institution}{dates}").font.size = Pt(9)

    # Certifications
    if cv.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in cv.certifications:
            date_str = f" ({cert.date})" if cert.date else ""
            p = doc.add_paragraph()
            run = p.add_run(cert.name)
            run.bold = True
            run.font.size = Pt(9)
            p.add_run(f" — {cert.issuer}{date_str}").font.size = Pt(9)

    # Languages
    if cv.languages:
        _add_section_heading(doc, "Languages")
        doc.add_paragraph(" · ".join(cv.languages))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def render_letter_docx(
    letter: GeneratedLetter, path: Path, *, profile: Profile | None = None
) -> None:
    """Render a GeneratedLetter to an ATS-optimized DOCX file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = _BODY_COLOR

    # Header with candidate info
    if profile:
        name_para = doc.add_heading(profile.name, level=1)
        for run in name_para.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = _DARK_BLUE

        contact_parts = [profile.email]
        if profile.phone:
            contact_parts.append(profile.phone)
        if profile.address and profile.address.city:
            loc = profile.address.city
            if profile.address.province_state:
                loc += f", {profile.address.province_state}"
            contact_parts.append(loc)
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        for run in contact_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = _MUTED_COLOR

        doc.add_paragraph("")  # spacer

    # Date
    date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    for run in date_para.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = _MUTED_COLOR

    # Greeting
    greeting_para = doc.add_paragraph()
    run = greeting_para.add_run(letter.greeting)
    run.bold = True
    run.font.color.rgb = _DARK_BLUE

    # Body paragraphs (hook / proof / close)
    for para_text in (letter.hook, letter.proof, letter.close):
        if para_text and para_text.strip():
            doc.add_paragraph(para_text.strip())

    # Closing
    doc.add_paragraph(letter.closing)

    # Signature
    sig_para = doc.add_paragraph()
    run = sig_para.add_run(letter.signature_name)
    run.bold = True
    run.font.color.rgb = _DARK_BLUE

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
